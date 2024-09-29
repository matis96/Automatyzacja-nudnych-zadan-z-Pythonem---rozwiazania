import docx
f = open("guests.txt", 'r')
lines = f.readlines()
doc = docx.Document('Invitations.docx')
for line in lines:
    doc.add_paragraph('Z prawdziwa przyjemnoscia firma', style='first_line')
    doc.add_paragraph(line.strip(), style='2nd_line')
    doc.add_paragraph('z siedziba przy ul. Nowej 123 zaprasza dnia', style='3rd_line')
    doc.add_paragraph('1-go kwietnia', style='4th_line')
    doc.add_paragraph('na godzine 19:00', style='5th_line')
    doc.add_page_break()
doc.save('Invitations.docx')
